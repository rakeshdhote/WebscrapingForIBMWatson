from bs4 import BeautifulSoup
import requests
from docx import Document

base_html = 'http://www.ontarioparks.com'
home_html = base_html + '/en'

#%% the function parses web sections
def www_section_parser(soup, key, value, elem):
    parks_temp1 = soup.find(attrs={key: value})
    parks_temp2 = parks_temp1.find_all(elem)
    return parks_temp2

#%% Extract parks name and url in dictionary
def parks_information(home_html):
    r = requests.get(home_html)
    soup = BeautifulSoup(r.content)

    temp_parks = www_section_parser(soup, 'class','btn-group','li')

    # Create list of park dictionaries
    parks_list = [] 
    for item in temp_parks:
        parks = {}
        parks['name'] = item.text
        parks['url_ext'] = item.find('a').get('href')
        parks_list.append(parks)
    return parks_list

#%% Extract each park information in the park list
def scrape_parks(parks_list):
    document = Document() # Open Word Document

    for i in range(0,len(parks_list)):

        parkname = parks_list[i]['name']
        park_url = base_html + parks_list[i]['url_ext']

        print parkname

        r = requests.get(park_url)
        soup = BeautifulSoup(r.content)

        ## Park Introduction
        temp_intro = www_section_parser(soup, 'id','tabs-introduction','li')

        park_intro = []
        for item in temp_intro:
            park_intro.append(item.text)

        ## Park Camping
        temp_camping = www_section_parser(soup, 'id','tabs-camping','p')

        park_camping = []
        for item in temp_camping:
            park_camping.append(item.text)

        ## Park - Things to Do
        temp_tod = www_section_parser(soup, 'id','tabs-thingstodo','p')

        park_thingstodo = []
        for item in temp_tod:
            park_thingstodo.append(item.text)

        ## Park - Amenities
        temp_amenities = www_section_parser(soup, 'id','tabs-amenities','p')

        park_amenities = []
        for item in temp_amenities:
            park_amenities.append(item.text)

        ### Populating Word Document
        document.add_heading(parkname, level=1)
        document.add_heading('Introduction', level=2)
        for i in park_intro:
            document.add_paragraph(i) 

        document.add_heading('Camping in '+parkname, level=2)
        for i in park_camping:
            document.add_paragraph(i)

        document.add_heading('Things to do in '+parkname, level=2)
        for i in park_thingstodo:
            document.add_paragraph(i)

        document.add_heading('Amenities in '+parkname, level=2)
        for i in park_amenities:
            document.add_paragraph(i)

        document.add_paragraph('          ')
    #    document.add_page_break()

    document.save('Camp_Grounds_in_Ontario.docx')

#%%
if __name__ == "__main__":
    parks_list = parks_information(home_html)
    scrape_parks(parks_list)
    print '## Open the document: Camp_Grounds_in_Ontario.docx ##'